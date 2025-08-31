import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from module.text_generation import call_groq_llama

def generate_resume(resume_data, api_key, max_tokens=512, top_p=0.95, temperature=0.7, output_file="resume.docx"):
    doc = Document()
    
    name = resume_data.get("name", "")
    email = resume_data.get("email", "")
    phone = resume_data.get("phone", "")
    skills = resume_data.get("skills", [])
    job_role = resume_data.get("job_role", "")
    location = resume_data.get("location", "")
    work_experience = resume_data.get("work_experience", {})
    projects = resume_data.get("projects", {})
    certifications = resume_data.get("certifications", {})

    prompt = f"""
        Write a single ATS-friendly 'About the Applicant' section in 30–40 words
        for the role: {job_role}. 
        Return only the text. Do not add titles, labels, or extra sentences.
    """
    output = call_groq_llama(prompt, api_key, max_tokens, top_p, temperature)

    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.left_margin = Cm(2)

    def add_divider(doc):
        p = doc.add_paragraph()
        p_element = p._element
        p_pr = p_element.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single') 
        bottom.set(qn('w:sz'), '12')        
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000') 
        p_borders.append(bottom)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        p_pr.append(p_borders)

   
    name_p = doc.add_paragraph()
    run = name_p.add_run(f"{name}")
    run.bold = True
    run.font.size = Pt(36)
    run.font.name = "Segoe UI"
    name_p.paragraph_format.space_after = Pt(0)

   
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(0)
    tab_stops = contact_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15))  

    contact_p.add_run(f"{email}    ").font.size = Pt(10.5)
    contact_p.add_run(f"+91{phone}    ").font.size = Pt(10.5)
    contact_p.add_run(f"{location}").font.size = Pt(10.5)

    add_divider(doc)

    about_p = doc.add_paragraph()
    about_p.paragraph_format.space_before = Pt(6)
    about_p.paragraph_format.space_after = Pt(6)
    about_p.add_run(output).font.size = Pt(11)

    add_divider(doc)


    skills_p = doc.add_paragraph()
    skills_p.paragraph_format.space_before = Pt(6)
    skills_p.paragraph_format.space_after = Pt(6)
    skills_p.add_run("SKILLS").bold = True

    cols = 3
    rows = (len(skills) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)

    idx = 0
    for r in range(rows):
        row_cells = table.rows[r].cells
        for c in range(cols):
            if idx < len(skills):
                run = row_cells[c].paragraphs[0].add_run(skills[idx])
                run.font.size = Pt(10)
                idx += 1

    tbl = table._tbl 
    for row in tbl.findall(".//w:tr", tbl.nsmap):
        for cell in row.findall(".//w:tc", tbl.nsmap):
            tcPr = cell.find("w:tcPr", tbl.nsmap)
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell.append(tcPr)
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "0")  
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "FFFFFF")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    add_divider(doc)

    work_p = doc.add_paragraph()
    work_p.paragraph_format.space_before = Pt(6)
    work_p.paragraph_format.space_after = Pt(6)
    work_p.add_run("PROFESSIONAL EXPERIENCE").bold = True

    for key, exp in work_experience.items():
        company_p = doc.add_paragraph()
        company_p.paragraph_format.left_indent = Cm(0.4)
        company_run = company_p.add_run(exp['company'])
        company_run.bold = True
        company_run.font.size = Pt(11)

        tab_stops = company_p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15))

        date_run = company_p.add_run(f"{exp['start']} - {exp['end']}")
        date_run.bold = True
        date_run.font.size = Pt(11)

        job_p = doc.add_paragraph()
        job_p.paragraph_format.left_indent = Cm(0.5)
        job_role_run = job_p.add_run(exp["role"])
        job_role_run.bold = True
        job_role_run.font.size = Pt(11)

        work_exp_prompt = f"""
        Write a single ATS-friendly 'Work Experience' section in 40–50 words
        for the role: {exp['role']} at {exp['company']}.
        Return only the text. Do not add titles, labels, or extra sentences.
        """
        description_output = call_groq_llama(work_exp_prompt, api_key, max_tokens, top_p, temperature)

        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_before = Pt(1)
        desc_p.paragraph_format.space_after = Pt(4)
        desc_p.paragraph_format.left_indent = Cm(0.5)
        desc_p.add_run(description_output.strip()).font.size = Pt(11)

    add_divider(doc)

    cert_p = doc.add_paragraph()
    cert_p.paragraph_format.space_before = Pt(6)
    cert_p.paragraph_format.space_after = Pt(6)
    cert_p.add_run("CERTIFICATIONS").bold = True

    for key, cert in certifications.items():
        c_p = doc.add_paragraph()
        c_p.paragraph_format.left_indent = Cm(0.5)
        c_run = c_p.add_run(f"{cert['name']}, {cert['organization']}")
        c_run.bold = True
        c_run.font.size = Pt(11)

    add_divider(doc)

    proj_heading = doc.add_paragraph()
    proj_heading.paragraph_format.space_before = Pt(6)
    proj_heading.paragraph_format.space_after = Pt(6)
    proj_heading.add_run("PROJECTS").bold = True

    for key, project in projects.items():
        title_p = doc.add_paragraph()
        title_p.paragraph_format.left_indent = Cm(0.4)
        title_run = title_p.add_run(project["title"])
        title_run.bold = True
        title_run.font.size = Pt(11)

        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent = Cm(0.4)
        desc_run = desc_p.add_run(project["description"])
        desc_run.font.size = Pt(11)

    doc.save(output_file)
    return output_file
